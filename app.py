import os
from fastapi import FastAPI, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.oxml import ns
from ebooklib import epub
from typing import List

from ai_translation import translate_text_pipeline, LANG_NAMES
from pdf_utils import generate_shaped_pdf

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

FONTS_DIR = "fonts"
os.makedirs(FONTS_DIR, exist_ok=True)

FONT_FILES = {
    "default": {
        "ttf": os.path.join(FONTS_DIR, "NotoSans-Regular.ttf"),
        "family": "Noto Sans",
    },
    "hi": {
        "ttf": os.path.join(FONTS_DIR, "NotoSansDevanagari-Regular.ttf"),
        "family": "Noto Sans Devanagari",
    },
    "es": {
        "ttf": os.path.join(FONTS_DIR, "NotoSans-Regular.ttf"),
        "family": "Noto Sans",
    },
    "de": {
        "ttf": os.path.join(FONTS_DIR, "NotoSans-Regular.ttf"),
        "family": "Noto Sans",
    },
}


def get_font(lang: str):
    return FONT_FILES.get(lang, FONT_FILES["default"])


@app.get("/health")
def health():
    return {"message": "AI Multilanguage Translation Studio Running"}


@app.post("/translate-file")
async def translate_file(file: UploadFile, languages: List[str] = Form(...)):
    raw = await file.read()

    text = None
    for enc in ["utf-8", "utf-8-sig", "latin-1"]:
        try:
            text = raw.decode(enc)
            break
        except:
            pass

    if not text:
        raise HTTPException(status_code=400, detail="Invalid file encoding")

    results = []

    for lang in languages:
        result = translate_text_pipeline(text, lang)
        translated = result["translated"]

        if not translated.strip():
            raise HTTPException(status_code=500, detail=f"Translation failed for {lang}")

        font = get_font(lang)
        language_name = LANG_NAMES.get(lang, lang)

        # ------------------------- PDF -------------------------
        pdf_file = f"{lang}_output.pdf"
        pdf_path = os.path.join(OUTPUT_DIR, pdf_file)

        generate_shaped_pdf(
            text=translated,
            pdf_path=pdf_path,
            font_path=font["ttf"],
            language_name=language_name,
        )

        # ------------------------- DOCX -------------------------
        docx_file = f"{lang}_output.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_file)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = font["family"]
        rFonts = style.element.rPr.rFonts
        rFonts.set(ns.qn("w:eastAsia"), font["family"])
        style.font.size = Pt(12)

        for line in translated.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = font["family"]
            run.font.size = Pt(12)

        doc.save(docx_path)

        # ------------------------- EPUB -------------------------
        epub_file = f"{lang}_output.epub"
        epub_path = os.path.join(OUTPUT_DIR, epub_file)

        book = epub.EpubBook()
        book.set_identifier(f"trans-{lang}")
        book.set_title(f"Translated – {language_name}")
        book.set_language(lang)
        book.add_author("AI Translation Studio")

        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        css = f"""
        body {{
            font-family: '{font['family']}';
            white-space: pre-wrap;
        }}
        """

        css_item = epub.EpubItem(
            uid="style",
            file_name="styles/style.css",
            media_type="text/css",
            content=css.encode("utf-8")
        )
        book.add_item(css_item)

        if os.path.exists(font["ttf"]):
            with open(font["ttf"], "rb") as f:
                font_bytes = f.read()
            book.add_item(
                epub.EpubItem(
                    uid="font",
                    file_name=f"fonts/{os.path.basename(font['ttf'])}",
                    media_type="application/x-font-ttf",
                    content=font_bytes
                )
            )

        html_lines = []
        for line in translated.split("\n"):
            safe = (
                line.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                or "&nbsp;"
            )
            html_lines.append(f"<p>{safe}</p>")

        chapter = epub.EpubHtml(title="Chapter 1", file_name="chap1.xhtml")
        chapter.content = f"""
        <html>
            <head><link rel="stylesheet" href="styles/style.css" /></head>
            <body>{''.join(html_lines)}</body>
        </html>
        """

        book.add_item(chapter)
        book.spine = ["nav", chapter]
        book.toc = (epub.Link("chap1.xhtml", "Chapter 1", "chap1"),)

        epub.write_epub(epub_path, book)

        results.append({
            "language": lang,
            "model": result["used_model"],
            "files": {
                "pdf": f"http://127.0.0.1:5000/download/{pdf_file}",
                "docx": f"http://127.0.0.1:5000/download/{docx_file}",
                "epub": f"http://127.0.0.1:5000/download/{epub_file}",
            }
        })

    return {"results": results}


@app.get("/download/{filename}")
def download(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path, filename=filename)
